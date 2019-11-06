#!/usr/bin/env python3
# encoding: utf-8
# -*- coding: utf-8 -*-
# @Time : 2018-12-22 19:53
# @Author : Leon Chu <ieno_chu@apple.com>
from docx import Document
from docx.shared import Inches, RGBColor
import docx
import os
from docx.oxml.ns import qn


def readdocx(file_name):
    res = ''
    document = Document(file_name)  # 打开文件demo.docx
    for paragraph in document.paragraphs:
        res += paragraph.text + '\n'
    return res


def write_docx(file_name, arrs, percentage=None):
    document = Document()
    document.styles['Normal'].font.name = u'微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    p = document.add_paragraph('')
    if percentage:
        p.add_run('------------------统计---------------------\n')
        run = p.add_run('小学:' + str(percentage['primary'][0]) + '% -' + str(percentage['primary'][1]) + '个\n')
        run.font.color.rgb = RGBColor.from_string('2F4F4F')
        run = p.add_run('初一:' + str(percentage['middle_one'][0]) + '% -' + str(percentage['middle_one'][1]) + '个\n')
        run.font.color.rgb = RGBColor.from_string('000066')
        run = p.add_run('初二、初三:' + str(percentage['middle'][0]) + '% -' + str(percentage['middle'][1]) + '个\n')
        run.font.color.rgb = RGBColor.from_string('4169E1')
        run = p.add_run('高中必修:' + str(percentage['must'][0]) + '% -' + str(percentage['must'][1]) + '个\n')
        run.font.color.rgb = RGBColor.from_string('DAA520')
        run = p.add_run('高中选修:' + str(percentage['select'][0]) + '% -' + str(percentage['select'][1]) + '个\n')
        run.font.color.rgb = RGBColor.from_string('3CB371')
        run = p.add_run('超纲:' + str(percentage['out'][0]) + '% -' + str(percentage['out'][1]) + '个\n')
        run.font.color.rgb = RGBColor.from_string('FF4500')
        p.add_run('-----------------------------------------\n\n')
    for line in arrs:
        print(line)
        run = p.add_run(line[0])
        if line[1][0] == '#':
            run.font.color.rgb = RGBColor.from_string(line[1][1:])
        elif line[1] == 's':
            p.add_run(' ')
        elif line[1] == 'n':
            p.add_run('\n')

    p.add_run("\n=======  单词表 =======\n")
    p.add_run("\n小学单词 \n")
    for line in arrs:
        if line[1] == '#2F4F4F':
            if len(line[0]) > 2:
                run = p.add_run(line[0] + " \n")
    p.add_run("\n初中单词\n")
    for line in arrs:
        if line[1] == '#4169E1':
            if len(line[0]) > 2:
                run = p.add_run(line[0] + " \n")
    p.add_run("\n高中必修单词\n")
    for line in arrs:
        if line[1] == '#DAA520':
            if len(line[0]) > 2:
                run = p.add_run(line[0] + " \n")
    p.add_run("\n高中选修单词\n")
    for line in arrs:
        if line[1] == '#3CB371':
            if len(line[0]) > 2:
                run = p.add_run(line[0] + " \n")
    print(arrs)
    document.save(file_name)
    return file_name



def write_export_docx(file_name, arrs):
    document = Document()
    document.styles['Normal'].font.name = u'微软雅黑'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    p = document.add_paragraph('')
    return file_name



if __name__ == '__main__':
    # test2()
    # print(readdocx('demp2.docx'))
    test = [['leon', '#3CB371'], ['lrc', '#00FF00']]
    write_docx('demp2.docx', test)
    # a1 = [1, 2, 3]
    # a2 = [4, 5, 6]
    # a3 = a1 + a2
    # print(a3)
    print(os.path.dirname('/abc/as.txt'))


